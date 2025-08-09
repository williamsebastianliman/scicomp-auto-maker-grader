import os
import io
import sys
import random as rd
import numpy as np
from PIL import Image, ImageDraw, ImageFont, ImageOps
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docxtpl import DocxTemplate, InlineImage
from docx.shared import Inches
from docx import Document
import zipfile
import subprocess
import tempfile
import copy
from lxml import etree
from docx.oxml import parse_xml
from docx.shared import Pt
import shutil

try:
    from win32com import client
    HAS_WORD = True
except Exception:
    client = None
    HAS_WORD = False

# Use platform‑appropriate default font
DEFAULT_FONT = (
    r"C:\\Windows\\Fonts\\consola.ttf"
    if os.name == "nt"
    else "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf"
)


def _render_math_png(text: str, out_path: str, fontsize: int = 28) -> None:
    """Render TeX‑like math (matplotlib mathtext) to a transparent PNG."""
    os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
    fig = plt.figure(figsize=(6, 1))
    fig.patch.set_alpha(0.0)
    plt.axis("off")
    # Wrap with $...$ to enable subscripts/superscripts and other math syntax
    plt.text(0.5, 0.5, f"${text}$", ha="center", va="center", fontsize=fontsize)
    fig.savefig(out_path, dpi=300, bbox_inches="tight", pad_inches=0.2, transparent=True)
    plt.close(fig)


def _render_text_png(text: str, out_path: str, fontsize: int = 28) -> None:
    """Render plain unicode text (e.g., Riemann integrals) to a PNG."""
    os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
    fig = plt.figure(figsize=(6, 1))
    fig.patch.set_alpha(0.0)
    plt.axis("off")
    plt.text(0.5, 0.5, text, ha="center", va="center", fontsize=fontsize)
    fig.savefig(out_path, dpi=300, bbox_inches="tight", pad_inches=0.2, transparent=True)
    plt.close(fig)


def _add_border_to_png(path: str, border_px: int = 3) -> None:
    """Add a thin black border around an image in‑place."""
    if not os.path.isfile(path):
        return
    im = Image.open(path).convert("RGBA")
    # Add a white margin and then draw a black rectangle around the new canvas
    im = ImageOps.expand(im, border=border_px, fill="white")
    draw = ImageDraw.Draw(im)
    w, h = im.size
    draw.rectangle([0, 0, w - 1, h - 1], outline="black", width=1)
    im.save(path)

def add_border(image_path: str, border_px: int = 12) -> None:
    """
    Add a bold black border to a PNG image.  This helper is used to
    reinforce the border thickness around figures when they are inserted
    into the report.  The border width can be adjusted via ``border_px``.
    """
    if not os.path.isfile(image_path):
        return
    img = Image.open(image_path).convert("RGB")
    img = ImageOps.expand(img, border=border_px, fill="black")
    img.save(image_path)

def iter_all_paragraphs(doc: Document):
    """
    Yield all paragraphs in the given Document, including those in tables,
    headers, and footers.  This allows token replacement to operate on
    content throughout the document, not just the main body.
    """
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for sect in doc.sections:
        for part in (sect.header, sect.footer):
            for p in part.paragraphs:
                yield p
            for tbl in part.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p

def latex_to_omml_xml(latex: str) -> str:
    """
    Convert a LaTeX math snippet to an OMML XML string using Pandoc.
    The snippet should contain only the mathematical expression (no
    surrounding dollar signs).  This function writes the LaTeX into a
    temporary file, invokes pandoc to convert it to a docx, and then
    extracts the first OMML object from the resulting document.
    """
    latex = latex.strip()
    # Ensure inline math delimiters for pandoc
    if not (latex.startswith("$") or latex.startswith("\\(")):
        latex = f"${latex}$"
    with tempfile.NamedTemporaryFile(suffix=".tex", mode="w", delete=False) as tf:
        tf.write(latex)
        src = tf.name
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as of:
        out_docx = of.name
    try:
        subprocess.run(
            ["pandoc", src, "-f", "latex", "-t", "docx", "-o", out_docx],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
        d = Document(out_docx)
        
        for p in iter_all_paragraphs(d):
            nodes = p._p.xpath(".//*[local-name()='oMath' or local-name()='oMathPara']")
            if nodes:
                return etree.tostring(nodes[0], encoding='unicode')

    finally:
        # Cleanup temporary files
        try:
            os.unlink(src)
        except Exception:
            pass
        try:
            os.unlink(out_docx)
        except Exception:
            pass
    return ""

def _replace_in_paragraph_once(p, placeholder: str, omml_xml: str) -> bool:
    """
    Replace a single occurrence of ``placeholder`` in paragraph ``p`` with
    the OMML object specified by ``omml_xml``.  The remainder of the
    paragraph text is preserved.  Returns True if replacement occurred.
    """
    full = "".join(r.text for r in p.runs)
    idx = full.find(placeholder)
    if idx == -1:
        return False
    left = full[:idx]
    right = full[idx + len(placeholder):]
    for r in p.runs:
        r.text = ""
    if left:
        p.add_run(left)
    run_math = p.add_run()
    run_math._r.append(parse_xml(omml_xml))
    if right:
        p.add_run(right)
    return True

def replace_placeholder_with_omml_all(doc: Document, placeholder: str, omml_xml: str,
                                      space_before_pt: int | None = None,
                                      space_after_pt: int | None = None) -> int:
    """
    Replace all occurrences of ``placeholder`` in the document with the
    provided OMML XML.  Adjust paragraph spacing for paragraphs where
    replacements occur using ``space_before_pt`` and ``space_after_pt``.
    Returns the number of replacements made.
    """
    count = 0
    for p in iter_all_paragraphs(doc):
        while placeholder in p.text:
            if not _replace_in_paragraph_once(p, placeholder, omml_xml):
                break
            pf = p.paragraph_format
            if space_before_pt is not None:
                pf.space_before = Pt(space_before_pt)
            if space_after_pt is not None:
                pf.space_after = Pt(space_after_pt)
            count += 1
    return count


def _replace_tokens_with_images(docx_path: str, mapping: dict) -> None:
    """
    Replace any paragraph containing one or more tokens with the corresponding
    images. Tokens are keys of ``mapping``, values are file paths to PNGs.

    We iterate through all paragraphs in the document. If a paragraph's text
    contains any token, we clear the paragraph and insert the images in the
    order they appear. Each image is inserted at a width of 4 inches (to
    approximate the sizing used in the original template). A newline is added
    after each image to maintain separation when multiple tokens are present.
    """
    if not mapping:
        return
    doc = Document(docx_path)
    for p in doc.paragraphs:
        hits = [t for t in mapping.keys() if t in p.text]
        if not hits:
            continue
        # Clear existing text in the paragraph
        try:
            p.clear()  # python‑docx >= 0.8.11
        except AttributeError:
            p.text = ""
        for token in hits:
            run = p.add_run()
            run.add_picture(mapping[token], width=Inches(4))
            # Add a newline after each image to separate multiple tokens
            p.add_run("\n")
    doc.save(docx_path)

# -----------------------------------------------------------------------------
# Helper functions for converting LaTeX math to OMML (Word equations) and
# injecting them into a docx. These functions allow us to preserve the
# template-driven layout of the report while rendering equations using
# Pandoc. The generated OMML fragments are inserted using python‑docx.

def _latex_to_omml_elements(latex: str, display: bool = True):
    """
    Use Pandoc to convert a single LaTeX math expression into OMML (Word
    equation) elements. Returns a list of OxmlElement ready to append into a
    w:p. If ``display`` is True, the expression is treated as a block
    equation; otherwise inline.
    """
    # Build a temporary markdown file containing the LaTeX expression
    md = f"$$\n{latex}\n$$\n" if display else f"${latex}$\n"
    with tempfile.TemporaryDirectory() as td:
        md_path = os.path.join(td, "eq.md")
        out_docx = os.path.join(td, "eq.docx")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(md)
        # Call pandoc to convert markdown to docx
        subprocess.run(
            ["pandoc", md_path, "-f", "markdown", "-t", "docx", "-o", out_docx],
            check=True,
        )
        # Extract the OMML elements from the generated docx
        with zipfile.ZipFile(out_docx) as zf:
            xml = zf.read("word/document.xml")
        root = etree.fromstring(xml)
        ns = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
        }
        omml_nodes = root.xpath("//*[local-name()='oMathPara' or local-name()='oMath']")
        elements = []
        for node in omml_nodes:
            frag = etree.tostring(node)
            elements.append(parse_xml(frag))
        return elements


def _replace_token_with_omml(docx_path: str, token_to_latex: dict) -> None:
    """
    Open a docx, find paragraphs containing tokens, and replace the paragraph
    contents with OMML equations generated from the corresponding LaTeX
    expressions. Multiple tokens in a paragraph are processed in order.
    """
    doc = Document(docx_path)
    for p in doc.paragraphs:
        text = p.text
        if not text:
            continue
        hits = [token for token in token_to_latex.keys() if token in text]
        if not hits:
            continue
        # Remove all existing runs
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        # Insert OMML elements
        for token in hits:
            latex = token_to_latex[token]
            # Determine display vs inline: treat all as display for clarity
            for el in _latex_to_omml_elements(latex, display=True):
                p._p.append(copy.deepcopy(el))
        # Optionally insert a line break after each equation
        p.add_run("\n")
    doc.save(docx_path)

# -----------------------------------------------------------------------------
# Helper functions for constructing LaTeX and Markdown reports and invoking
# Pandoc. These functions allow us to build a report in a deterministic,
# human‑readable format and convert it into a .docx with native Office Math
# equations. They are used in place of the earlier docxtpl/Word COM pipeline.

def _coeffs_to_latex(coeffs, name="f"):
    """
    Convert a list of polynomial coefficients into a LaTeX equation string.
    ``coeffs`` should be ordered by increasing power (c[0] is constant term).
    The resulting string looks like f(x) = ax^n + bx^{n-1} + ... .
    """
    terms = []
    deg = len(coeffs) - 1
    for e in range(deg, -1, -1):
        c = coeffs[e]
        if c == 0:
            continue
        s = "-" if c < 0 else "+" if terms else ""
        a = abs(c)
        if e == 0:
            core = f"{a}"
        elif e == 1:
            core = "x" if a == 1 else f"{a}x"
        else:
            core = f"x^{e}" if a == 1 else f"{a}x^{e}"
        terms.append(f"{s} {core}".strip())
    poly = " ".join(terms) if terms else "0"
    return f"{name}(x) = {poly}"


def _system_to_latex(A, b):
    """
    Convert a system of linear equations (matrix A and vector b) into a
    LaTeX aligned environment. The result is suitable for embedding in
    display math mode. Each equation is rendered on its own line with
    subscripts for variables.
    """
    lines = []
    for row, rhs in zip(A, b):
        parts = []
        for i, coeff in enumerate(row, start=1):
            s = "-" if coeff < 0 else "+" if parts else ""
            a = abs(coeff)
            term = f"x_{i}" if a == 1 else f"{a}x_{i}"
            parts.append(f"{s} {term}".strip())
        lhs = " ".join(parts) if parts else "0"
        lines.append(f"{lhs} = {rhs}")
    body = " \\ \n".join(lines)
    return f"\\begin{{aligned}}\n{body}\n\\end{{aligned}}"


def _write_markdown(soe, nr, rp, md_path):
    """
    Write a markdown report combining images and LaTeX equations. The images
    generated by the payloads are referenced relative to the markdown file.
    ``soe``: gauss_seidel_payload instance; ``nr``: newton_rapshon_payload;
    ``rp``: rimeann_payload; ``md_path``: path to write the markdown file.
    The markdown includes headings, images, and LaTeX math delimited by $$.
    """
    # Gauss–Seidel: use first few systems for brevity
    gs_blocks = []
    for system in soe.soe_list[:3]:
        gs_blocks.append(_system_to_latex(system.x_matrix, system.y_matrix))
    gs_math = "\n\n".join(f"$$\n{blk}\n$$" for blk in gs_blocks)

    # Newton–Raphson: function string in LaTeX
    nr_func = _coeffs_to_latex(nr.coeffs, name="f")

    # Riemann integral: build polynomial from rp.coeffs (index is power)
    r_terms = []
    for k, c in enumerate(rp.coeffs):
        if c == 0:
            continue
        s = "-" if c < 0 else "+" if r_terms else ""
        a = abs(c)
        if k == 0:
            core = f"{a}"
        elif k == 1:
            core = "x" if a == 1 else f"{a}x"
        else:
            core = f"x^{k}" if a == 1 else f"{a}x^{k}"
        r_terms.append(f"{s} {core}".strip())
    r_poly = " ".join(r_terms) if r_terms else "0"
    riemann_tex = f"\\int_{{{rp.a}}}^{{{rp.b}}} \\left({r_poly}\\right)\\,dx"

    md = f"""# Report\n\n"""
    md += "## Gauss–Seidel\n\n"
    md += "![Gauss–Seidel Iterations](../images/gauss_seidel_output.png)\n\n"
    md += "### Systems (editable equations)\n\n"
    md += gs_math + "\n\n"
    md += "## Regression\n\n"
    md += "![Best Fit](../images/best_fit.png)\n\n"
    md += "## Newton–Raphson\n\n"
    md += f"**Function:**  ${nr_func}$\n\n"
    md += "![Newton–Raphson Iterations](../images/newton_rapshon_image.png)\n\n"
    md += "## Riemann\n\n"
    md += f"**Integral:**  $$ {riemann_tex} $$\n"
    # Write out the markdown file
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md)


def _run_pandoc(md_path, out_docx, reference_doc=None):
    """
    Convert a markdown file to a .docx using Pandoc. If ``reference_doc`` is
    provided and exists, it is passed to pandoc via --reference-doc to
    apply custom styles. Assumes ``pandoc`` is installed and available in
    the environment.
    """
    import subprocess
    cmd = ["pandoc", md_path, "-o", out_docx]
    if reference_doc and os.path.isfile(reference_doc):
        cmd += ["--reference-doc", reference_doc]
    # Ensure resource path includes the report directory and the project root so
    # images located in ./images are discoverable by pandoc. The paths are
    # separated by the OS path separator (':' on Unix, ';' on Windows).
    report_dir = os.path.dirname(md_path)
    root_dir = os.getcwd()
    resource_paths = os.pathsep.join([report_dir, root_dir])
    cmd += ["--resource-path", resource_paths]
    subprocess.run(cmd, check=True)

def save_output_as_png(
    output_image: str = "output.png",
    font_path: str = DEFAULT_FONT,
    font_size: int = 36,
    padding: int = 10,
    dpi: tuple = (900, 900),
    line_spacing: float = 1.5,
) -> callable:
    """
    Decorator that captures printed output from a function and writes it into
    an image. The resulting PNG is saved to ``output_image`` at the specified
    DPI. If the output directory does not exist, it is created.
    """
    def decorator(fn):
        def wrapper(*args, **kwargs):
            # Redirect stdout into a buffer
            old_stdout = sys.stdout
            buf = io.StringIO()
            sys.stdout = buf
            try:
                result = fn(*args, **kwargs)
            finally:
                sys.stdout = old_stdout
            lines = buf.getvalue().splitlines() or [""]
            # Ensure the output directory exists
            os.makedirs(os.path.dirname(os.path.abspath(output_image)), exist_ok=True)
            font = ImageFont.truetype(font_path, font_size)
            dummy = Image.new("RGB", (1, 1))
            draw = ImageDraw.Draw(dummy)
            bboxes = [draw.textbbox((0, 0), line, font=font) for line in lines]
            max_w = max(x1 - x0 for x0, y0, x1, y1 in bboxes)
            line_h = max(y1 - y0 for x0, y0, x1, y1 in bboxes)
            step_h = int(line_h * line_spacing)
            img_w = max_w + padding * 2
            img_h = step_h * len(lines) + padding * 2
            img = Image.new("RGB", (img_w, img_h), "white")
            draw = ImageDraw.Draw(img)
            y = padding
            for line in lines:
                draw.text((padding, y), line, font=font, fill="black")
                y += step_h
            img.save(output_image, dpi=dpi)
            # Print for debugging
            print(f"Saved output to {output_image}")
            return result
        return wrapper
    return decorator

def rand_except(min_val: int, max_val: int, except_val: int) -> int:
    """Return a random integer between min_val and max_val inclusive, excluding except_val."""
    while True:
        num = rd.randint(min_val, max_val)
        if num != except_val:
            return num

def rand_interval(min1: int, max1: int, min2: int, max2: int) -> int:
    """Return a random integer from either of two intervals with equal probability."""
    return rd.randint(min1, max1) if rd.randint(0, 1) == 0 else rd.randint(min2, max2)

def rand_index(arr_size: int, selected_amount: int) -> list:
    """Return a list of unique random indices from the range [0, arr_size)."""
    selected_indices = []
    idx_list = list(range(arr_size))
    for _ in range(selected_amount):
        max_idx = len(idx_list) - 1
        selected_idx = rd.randint(0, max_idx)
        num = idx_list[selected_idx]
        selected_indices.append(num)
        idx_list.pop(selected_idx)
    return selected_indices

class gauss_seidel_equation_system:
    """
    Generate a single system of linear equations intended for the Gauss–Seidel method.
    Each system consists of three equations. Systems can be optionally diagonal
    dominant, which increases the chance of convergence.
    """
    def is_diagonal_domminant(self, matrix):
        matrix_abs = np.array(np.abs(matrix))
        matrix_diag = np.array(np.diag(matrix_abs))
        np.fill_diagonal(matrix_abs, 0)
        off_diag = np.sum(matrix_abs, axis=1)
        return np.all(matrix_diag > off_diag)

    def gauss_seidel(self, x, y):
        if not self.is_diagonal_domminant(x):
            return [0, 0, 0], False
        x = np.array(x, dtype=float)
        y = np.array(y, dtype=float)
        max_iter = 100
        threshold = 0.034
        x_diag = np.array(np.diag(x))
        np.fill_diagonal(x, 0)
        res = np.zeros(len(x_diag))
        for _ in range(max_iter):
            res_new = np.array(res)
            for idx, row in enumerate(x):
                res_new[idx] = (y[idx] - np.dot(row, res_new)) / x_diag[idx]
            if np.linalg.norm(res_new - res) < threshold:
                return res_new, True
            res = res_new
        return res, False

    def __init__(self, coeff_amount: int, equation_amount: int, is_diagonally_dominant: bool):
        x_matrix = []
        y_matrix = []
        for i in range(equation_amount):
            coeff_list = [0, 0, 0]
            diagonal_dominant_coeff = rand_interval(-99, -50, 50, 99)
            diagonal_dominant_coeff += 10 if diagonal_dominant_coeff >= 0 else -10
            coeff_list[i] = diagonal_dominant_coeff
            offset = (
                rd.randint(-30, -20) - 15
                if is_diagonally_dominant
                else rd.randint(20, 30) + 15
            )
            splitter = rd.randint(1, 9)
            total_split = abs(diagonal_dominant_coeff) + offset
            split1 = (total_split * splitter) // 10
            split2 = (total_split * (10 - splitter)) // 10
            if split1 == 0:
                split1 += rand_except(-3, 3, 0)
            elif split2 == 0:
                split2 += rand_except(-3, 3, 0)
            split_list = [split1, split2]
            for j in range(coeff_amount):
                if i != j:
                    coeff_list[j] = split_list.pop(0)
            b = rd.randint(-999, 999)
            x_matrix.append(coeff_list)
            y_matrix.append(b)
        self.x_matrix = x_matrix
        self.y_matrix = y_matrix
        self.coeff_amount = coeff_amount
        self.equation_amount = equation_amount
        self.res = self.gauss_seidel(self.x_matrix, self.y_matrix)

    def debug(self, var_list: list = ["x_1", "x_2", "x_3"]) -> str:
        system_eq = ""
        for eq_idx, eq in enumerate(self.x_matrix):
            eq_str = ""
            for idx, coeff in enumerate(eq):
                if idx == 0:
                    eq_str += f"{coeff}{var_list[idx]}"
                else:
                    eq_str += (f"+{coeff}{var_list[idx]}" if coeff >= 0 else f"{coeff}{var_list[idx]}")
            eq_str += f"={self.y_matrix[eq_idx]}"
            system_eq += eq_str
            if eq_idx < self.equation_amount - 1:
                system_eq += "\n"
        return system_eq

    def debug_answer(self):
        return self.res

class gauss_seidel_payload:
    """
    Build multiple Gauss–Seidel systems, print iterations for a few of them, save
    the printed output to an image, and build the x/y matrix payload for docxtpl.
    """
    def is_diagonal_domminant(self, matrix):
        matrix_abs = np.array(np.abs(matrix))
        matrix_diag = np.array(np.diag(matrix_abs))
        np.fill_diagonal(matrix_abs, 0)
        off_diag = np.sum(matrix_abs, axis=1)
        return np.all(matrix_diag > off_diag)

    @save_output_as_png(
        output_image="./images/gauss_seidel_output.png",
        font_size=36,
        line_spacing=1.5,
    )
    def gauss_seidel_print(self, x_system, y_system, limit=3):
        count = 0
        for idx, (x, y) in enumerate(zip(x_system, y_system)):
            count += 1
            if count > limit:
                return
            print(f"System of Equation {idx}")
            print("=================================")
            if not self.is_diagonal_domminant(x):
                print("Not Diagonally Dominant!")
                continue
            print("Diagonally Dominant!")
            x = np.array(x, dtype=float)
            y = np.array(y, dtype=float)
            max_iter = 100
            threshold = 0.099
            x_diag = np.diag(x).astype(float)
            np.fill_diagonal(x, 0.0)
            res = np.zeros(len(x_diag), dtype=float)
            for i in range(max_iter):
                res_new = res.copy()
                for idx, row in enumerate(x):
                    res_new[idx] = (y[idx] - np.dot(row, res_new)) / x_diag[idx]
                diff = np.linalg.norm(res_new - res)
                print("Iteration:", i)
                print(f"Result: {res_new}")
                if diff < threshold:
                    print(f"Convergence: {res_new}\n")
                    break
                res = res_new
            else:
                print("Not Convergence (Max Iter Reached)\n")

    def xy_builder(self, x_list, y_list) -> str:
        lines = []
        lines.append("x = [")
        for mi, mat in enumerate(x_list):
            mat_comma = "," if mi < len(x_list) - 1 else ""
            lines.append("    [")
            for ri, row in enumerate(mat):
                row_comma = "," if ri < len(mat) - 1 else ""
                row_str = ", ".join(str(v) for v in row)
                lines.append(f"        [{row_str}]{row_comma}")
            lines.append(f"    ]{mat_comma}")
        lines.append("]")
        lines.append("")
        lines.append("y = [")
        for ri, row in enumerate(y_list):
            row_comma = "," if ri < len(y_list) - 1 else ""
            row_str = ", ".join(str(v) for v in row)
            lines.append(f"    [{row_str}]{row_comma}")
        lines.append("]")
        return "\n".join(lines)

    def __init__(self, system_of_eq_amount: int = 5):
        max_iter = rd.randint(50, 120)
        self.max_iter = str(max_iter)
        self.soe_list = []
        not_diagonal_amount = system_of_eq_amount // 3 or 1
        indices = rand_index(system_of_eq_amount, not_diagonal_amount)
        for i in range(system_of_eq_amount):
            eqs = gauss_seidel_equation_system(3, 3, is_diagonally_dominant=(i not in indices))
            self.soe_list.append(eqs)
        system_of_x = [soe.x_matrix for soe in self.soe_list]
        system_of_y = [soe.y_matrix for soe in self.soe_list]
        self.gauss_seidel_print(system_of_x, system_of_y)
        self.output = "./images/gauss_seidel_output.png"
        self.payload = self.xy_builder(system_of_x, system_of_y)

class regression_payload:
    """
    Create a random linear regression problem, compute the least squares fit,
    plot the points and the best‑fit line, and build a payload file for docxtpl.
    """
    def regression_solver(self, a, b):
        x = np.array(a)
        y = np.array(b)
        A = np.vstack([x, np.ones(len(x))]).T
        inv = np.linalg.inv(np.dot(A.T, A))
        inv = np.dot(inv, A.T)
        inv = np.dot(inv, y)
        return inv

    def __init__(self):
        m_mult = rd.randint(1, 9)
        b_mult = rd.randint(1, 9)
        m = rd.randint(2, 9)
        b = rd.randint(5, 25)
        m_fin = m * m_mult
        b_fin = b * b_mult
        x_points = np.linspace(1, 50, 50)
        x_points = np.array(x_points) * rd.randint(1, 5)
        x_point_fin = []
        y_point_fin = []
        for x in x_points:
            offset = rd.randint(-5, 5) / 10.0
            x_point_fin.append(round(float(x + offset), 3))
        for x in x_point_fin:
            offset_y = rd.randint(-10, 10)
            y = m_fin * x + b_fin
            y_real = y + ((offset_y / 100) * y) + offset_y
            y_point_fin.append(round(float(y_real), 3))
        self.x_point_fin = x_point_fin
        self.y_point_fin = y_point_fin
        m_fit, b_fit = self.regression_solver(x_point_fin, y_point_fin)
        self.m = m_fit
        self.b = b_fit
        self.image_url = self.plot()
        self.payload = self.payload_builder()

    def payload_builder(self, out_path="./result/2-matrix.txt") -> str:
        def fmt(v):
            s = f"{float(v):.3f}".rstrip("0").rstrip(".")
            return s if "." in s else s + ".0"
        def block(label, data, per_line=10):
            lines = []
            for i in range(0, len(data), per_line):
                chunk = ", ".join(fmt(v) for v in data[i : i + per_line])
                lines.append(f"    {chunk}")
            inner = ",\n".join(lines)
            return f"{label} = [\n{inner}\n]\n"
        x_block = block("x", self.x_point_fin)
        y_block = block("y", self.y_point_fin)
        payload = f"{x_block}\n{y_block}"
        os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(payload)
        return payload

    def plot(self):
        os.makedirs("./images", exist_ok=True)
        fig, ax = plt.subplots()
        ax.plot(self.x_point_fin, self.y_point_fin, "or", label="Coordinate")
        x_point_fin = np.array(self.x_point_fin)
        ax.plot(self.x_point_fin, self.m * x_point_fin + self.b, "b-", label="Best Fit")
        ax.legend()
        fig.savefig("./images/best_fit.png", dpi=900, bbox_inches="tight")
        plt.close(fig)
        return "./images/best_fit.png"

class newton_rapshon_payload:
    """
    Generate a polynomial and use the Newton–Raphson method to find a root. Print
    the iteration steps into an image and provide a function string for docxtpl.
    """
    def __init__(self, degree=5, tol=0.099, max_tries=200):
        self.degree = degree
        self.tol = tol
        self.max_iter = rd.randint(50, 100)
        for _ in range(max_tries):
            self.coeffs = self._gen_coeffs(degree)
            self.x0 = self.pick_x0()
            self.function = self.build_function(self.coeffs, "f")
            self.root = self.solver(tol=self.tol, max_iter=100)
            if getattr(self, "converged", False) and getattr(self, "iters", 0) <= 12:
                break
        self.solver_printed(tol=self.tol)

    def _gen_coeffs(self, degree):
        c = [0] * (degree + 1)
        c[degree] = rd.choice([-6, -5, -4, -3, -2, -1, 1, 2, 3, 4, 5, 6])
        max_const = 99
        min_high = 12
        beta = 1.3
        for p in range(degree):
            frac = (degree - p) / degree if degree > 0 else 1.0
            lim = max(min_high, int(round(1 + (max_const - 1) * (frac ** beta))))
            v = rd.randint(-lim, lim)
            if p > 0 and v == 0:
                v = rd.choice([-1, 1])
            c[p] = v
        return c

    def f(self, x):
        return sum(self.coeffs[k] * (x ** k) for k in range(self.degree + 1))

    def df(self, x):
        return sum(k * self.coeffs[k] * (x ** (k - 1)) for k in range(1, self.degree + 1))

    def pick_x0(self, span=(-10, 10), eps=1e-3, tries=200):
        lo = int(np.ceil(span[0]))
        hi = int(np.floor(span[1]))
        if lo > hi:
            lo, hi = hi, lo
        cand = list(range(lo, hi + 1))
        rd.shuffle(cand)
        for x0 in cand[:tries]:
            if abs(self.df(x0)) > eps:
                return x0
        return max(cand, key=lambda t: abs(self.df(t))) if cand else 0

    def _sup(self, n: int) -> str:
        return str(n).translate(str.maketrans("0123456789-", "⁰¹²³⁴⁵⁶⁷⁸⁹⁻"))

    def build_function(self, coeff_arr, name="f") -> str:
        terms = []
        deg = len(coeff_arr) - 1
        for e in range(deg, -1, -1):
            c = coeff_arr[e]
            if c == 0:
                continue
            a = abs(c)
            if e == 0:
                core = f"{a}"
            elif e == 1:
                core = "x" if a == 1 else f"{a}x"
            else:
                core = f"x{self._sup(e)}" if a == 1 else f"{a}x{self._sup(e)}"
            if not terms:
                terms.append(core if c > 0 else f"-{core}")
            else:
                terms.append((" + " if c > 0 else " - ") + core)
        poly = "".join(terms) if terms else "0"
        return f"{name}(x) = {poly}"

    def solver(self, tol=0.099, max_iter=100):
        x = float(self.x0)
        hist = []
        for i in range(max_iter):
            fx = self.f(x)
            dfx = self.df(x)
            if dfx == 0:
                self.root = None
                self.iters = i
                self.converged = False
                self.history = hist
                return None
            x_new = x - fx / dfx
            hist.append((i, x, fx, dfx, abs(x_new - x)))
            if abs(x_new - x) < tol:
                self.root = x_new
                self.iters = i + 1
                self.converged = True
                self.history = hist
                return x_new
            x = x_new
        self.root = x
        self.iters = len(hist)
        self.converged = False
        self.history = hist
        return x

    @save_output_as_png(
        output_image="./images/newton_rapshon_image.png",
        font_size=36,
        line_spacing=1.5,
    )
    def solver_printed(self, tol=0.099, fmt="{:.15f}"):
        x = float(self.x0)
        for k in range(1, self.max_iter + 1):
            fx = self.f(x)
            dfx = self.df(x)
            if dfx == 0:
                print("Derivative is zero; Newton step undefined.")
                self.converged = False
                return None
            x_new = x - fx / dfx
            print(f"Iteration {k}: {fmt.format(x_new)}")
            if abs(x_new - x) < tol:
                print(f"Root is found at x:  {fmt.format(x_new)}")
                self.root = x_new
                self.iters = k
                self.converged = True
                return x_new
            x = x_new
        print(f"Reached max iterations. Last x: {fmt.format(x)}")
        self.root = x
        self.iters = self.max_iter
        self.converged = False
        return x

class rimeann_payload:
    """
    Build a random polynomial integral problem and approximate it using
    multiple Riemann sum methods. Provide errors and string representations
    for inclusion in docxtpl.
    """
    def idx_to_method(self, idx):
        return {
            1: "Left Riemann",
            2: "Right Riemann",
            3: "Mid Riemann",
            4: "Trapezoid Riemann",
        }[idx]

    def __init__(self):
        i1, i2 = rd.sample([1, 2, 3, 4], 2)
        self.method1 = self.idx_to_method(i1)
        self.method2 = self.idx_to_method(i2)
        self._gen_problem()
        self.results = self.solve_all()
        self.exact_value = self.exact_integral()
        self.value1 = self.results[self.method1]
        self.value2 = self.results[self.method2]
        self.errors = {k: abs(v - self.exact_value) for k, v in self.results.items()}

    def _gen_problem(self):
        def rand_nz(lo, hi):
            v = 0
            while v == 0:
                v = rd.randint(lo, hi)
            return v
        self.degree = 3
        if self.degree == 1:
            c1 = rd.choice([-5, -4, -3, -2, -1, 1, 2, 3, 4, 5])
            c0 = rand_nz(-8, 8)
            self.coeffs = [c0, c1]
        elif self.degree == 2:
            c2 = rd.choice([-2, -1, 1, 2])
            c1 = rand_nz(-5, 5)
            c0 = rand_nz(-8, 8)
            self.coeffs = [c0, c1, c2]
        else:
            c3 = rd.choice([-1, 1])
            c2 = rand_nz(-3, 3)
            c1 = rand_nz(-5, 5)
            c0 = rand_nz(-8, 8)
            self.coeffs = [c0, c1, c2, c3]
        a = rd.randint(-4, 1)
        b = a + rd.randint(2, 5)
        self.a, self.b = a, b
        self.n = rd.randint(100, 400)
        self.function_str = self._poly_pretty(self.coeffs, "f")
        self.word_integral_str = self._word_integral()

    def _sup(self, n: int) -> str:
        return str(n).translate(str.maketrans("0123456789-", "⁰¹²³⁴⁵⁶⁷⁸⁹⁻"))

    def _poly_pretty(self, coeffs, name="f"):
        terms = []
        d = len(coeffs) - 1
        for e in range(d, -1, -1):
            c = coeffs[e]
            if c == 0:
                continue
            a = abs(c)
            if e == 0:
                core = f"{a}"
            elif e == 1:
                core = "x" if a == 1 else f"{a}x"
            else:
                core = f"x{self._sup(e)}" if a == 1 else f"{a}x{self._sup(e)}"
            s = core if c > 0 else f"-{core}"
            if terms and c > 0:
                s = " + " + s
            elif terms and c < 0:
                s = " - " + core
            terms.append(s)
        poly = "".join(terms) if terms else "0"
        return f"{name}(x) = {poly}"

    def _poly_linear(self):
        terms = []
        d = len(self.coeffs) - 1
        for e in range(d, -1, -1):
            c = self.coeffs[e]
            if c == 0:
                continue
            sign = (
                "+"
                if (c > 0 and terms)
                else "-"
                if (c < 0 and terms)
                else "-"
                if c < 0
                else ""
            )
            a = abs(c)
            if e == 0:
                core = f"{a}"
            elif e == 1:
                core = "x" if a == 1 else f"{a} x"
            else:
                core = f"x^{e}" if a == 1 else f"{a} x^{e}"
            terms.append(f"{sign} {core}".strip())
        return " ".join(terms) if terms else "0"

    def _word_integral(self):
        integral = "\u222B"
        return f"{integral}_{self.a}^{self.b} ({self._poly_linear()}) d x"

    def latex_integral_str(self) -> str:
        """
        Construct a LaTeX string representing the definite integral of the
        polynomial defined by ``self.coeffs`` over [a, b].  The polynomial
        terms are assembled in descending order of power with appropriate
        signs.  The result can be converted to an OMML equation via pandoc.
        """
        terms = []
        for e in range(len(self.coeffs) - 1, -1, -1):
            c = self.coeffs[e]
            if c == 0:
                continue
            sign = "+" if (c > 0 and terms) else "-" if (c < 0 and terms) else "-" if c < 0 else ""
            a = abs(c)
            if e == 0:
                core = f"{a}"
            elif e == 1:
                core = "x" if a == 1 else f"{a}x"
            else:
                core = f"x^{e}" if a == 1 else f"{a}x^{e}"
            terms.append(f"{sign} {core}".strip())
        poly = " ".join(terms) if terms else "0"
        return fr"\int_{{{self.a}}}^{{{self.b}}} \left({poly}\right)\,dx"

    def f_val(self, x: float) -> float:
        return sum(c * (x ** k) for k, c in enumerate(self.coeffs))

    def riemann(self, method: str, n: int = None) -> float:
        if n is None:
            n = self.n
        a, b = float(self.a), float(self.b)
        dx = (b - a) / n
        m = method.lower()
        if "left" in m:
            xs = (a + i * dx for i in range(n))
            return dx * sum(self.f_val(x) for x in xs)
        if "right" in m:
            xs = (a + i * dx for i in range(1, n + 1))
            return dx * sum(self.f_val(x) for x in xs)
        if "mid" in m:
            xs = (a + (i + 0.5) * dx for i in range(n))
            return dx * sum(self.f_val(x) for x in xs)
        if "trap" in m:
            s = 0.5 * (self.f_val(a) + self.f_val(b))
            s += sum(self.f_val(a + i * dx) for i in range(1, n))
            return dx * s
        raise ValueError(f"Unknown method: {method}")

    def solve_all(self, n: int = None) -> dict:
        if n is None:
            n = self.n
        return {
            "Left Riemann": self.riemann("left", n),
            "Right Riemann": self.riemann("right", n),
            "Mid Riemann": self.riemann("mid", n),
            "Trapezoid Riemann": self.riemann("trapezoid", n),
        }

    def exact_integral(self) -> float:
        a, b = float(self.a), float(self.b)
        total = 0.0
        for k, c in enumerate(self.coeffs):
            total += c * (b ** (k + 1) - a ** (k + 1)) / (k + 1)
        return total

def generate_docs(seed: int = None):
    """
    Generate images and a docx report. If a seed is provided, the random
    state is reset to make the output deterministic.  We rely on docxtpl
    to render the report according to the provided Word template and
    then replace equation placeholders with native OMML equations using
    Pandoc.  This preserves the template layout without requiring Word.
    """
    # Reset RNG
    if seed is not None:
        rd.seed(int(seed))
        np.random.seed(int(seed) & 0xFFFFFFFF)

    soe = gauss_seidel_payload()
    regression_payload()
    nr = newton_rapshon_payload()
    rp = rimeann_payload()

    for img_path in (
        "./images/gauss_seidel_output.png",
        "./images/best_fit.png",
        "./images/newton_rapshon_image.png",
    ):
        _add_border_to_png(img_path, border_px=3)

    tpl = DocxTemplate(os.path.abspath("./template/template.docx"))
    tpl.render(
        {
            "initial_value_gauss_seidel": "0",
            "max_iteration_gauss_seidel": soe.max_iter,
            "xy_def_gauss_seidel": soe.payload,
            "gauss_seidel_image": InlineImage(
                tpl, os.path.abspath("./images/gauss_seidel_output.png"), width=Inches(4)
            ),
            "regression_image": InlineImage(
                tpl, os.path.abspath("./images/best_fit.png"), width=Inches(4)
            ),
            "equation_newton_rapshon": nr.function,
            "first_guess_newton_rapshon": nr.x0,
            "iteration_newton_rapshon": nr.max_iter,
            "newton_rapshon_image": InlineImage(
                tpl, os.path.abspath("./images/newton_rapshon_image.png"), width=Inches(4)
            ),
            "method1_riemann": rp.method1,
            "method2_riemann": rp.method2,
            "point_riemann": rp.n,
        }
    )

    out_dir = os.path.abspath("./result")
    os.makedirs(out_dir, exist_ok=True)

    temp_path = os.path.join(out_dir, "temp.docx")
    tpl.save(temp_path)

    token_to_latex = {}
    for sys_idx, system in enumerate(soe.soe_list, start=1):
        for eq_idx, eq_text in enumerate(system.debug().split("\n"), start=1):
            token = f"EQ{sys_idx}{eq_idx}"
            token_to_latex[token] = eq_text
    if hasattr(rp, "latex_integral_str"):
        token_to_latex["RIEMANN_EQ"] = rp.latex_integral_str()
    else:
        token_to_latex["RIEMANN_EQ"] = rp.word_integral_str

    final_path = os.path.join(out_dir, "result.docx")
    shutil.copy2(temp_path, final_path)

    for img_path in (
        "./images/gauss_seidel_output.png",
        "./images/best_fit.png",
        "./images/newton_rapshon_image.png",
    ):
        add_border(img_path, border_px=12)

    # Open the temporary doc and inject OMML equations in place of tokens
    doc = Document(temp_path)
    for sys_idx, system in enumerate(soe.soe_list, start=1):
        eqs = system.debug().split("\n")
        n_eq = len(eqs)
        for eq_idx, eq_text in enumerate(eqs, start=1):
            token = f"EQ{sys_idx}{eq_idx}"
            omml_xml = latex_to_omml_xml(eq_text)
            after = 12 if eq_idx == n_eq else 0
            replace_placeholder_with_omml_all(
                doc, token, omml_xml,
                space_before_pt=0,
                space_after_pt=after
            )
    # Replace the Riemann placeholder with a single OMML integral.  No extra spacing needed.
    riemann_xml = latex_to_omml_xml(rp.latex_integral_str())
    replace_placeholder_with_omml_all(
        doc, "RIEMANN_EQ", riemann_xml,
        space_before_pt=0,
        space_after_pt=0
    )
    # Save the final document
    doc.save(final_path)
    return final_path

def clean_outputs(images_dir="./images", result_dir="./result"):
    """Remove and recreate the images and result directories."""
    for p in (images_dir, result_dir):
        if os.path.isdir(p):
            shutil.rmtree(p, ignore_errors=True)
        os.makedirs(p, exist_ok=True)

def make_cases(
    docs_count: int,
    out_root: str = "./output",
    src_doc: str = "./result/result.docx",
    src_txt: str = "./result/2-matrix.txt",
    seed: int = -1,
) -> list:
    """
    Generate multiple cases. For each case, generate the report with a deterministic
    seed (if provided) and copy the report and matrix file into a unique directory.
    """
    clean_outputs()
    os.makedirs(out_root, exist_ok=True)
    used_ids = set()
    cases = []
    for _ in range(docs_count):
        if seed != -1:
            case_seed = int(seed)
            case_id = f"{case_seed:07d}"
        else:
            rng = rd.Random()
            case_id = f"{rng.randint(0, 9_999_999):07d}"
            case_seed = int(case_id)
        case_dir = os.path.join(out_root, f"Case_{case_id}")
        while case_id in used_ids or os.path.exists(case_dir):
            rng = rd.Random()
            case_id = f"{rng.randint(0, 9_999_999):07d}"
            case_seed = int(case_id)
            case_dir = os.path.join(out_root, f"Case_{case_id}")
        generate_docs(seed=case_seed)
        if not os.path.isfile(src_doc):
            raise FileNotFoundError(f"Missing: {src_doc}")
        if not os.path.isfile(src_txt):
            raise FileNotFoundError(f"Missing: {src_txt}")
        os.makedirs(case_dir, exist_ok=False)
        used_ids.add(case_id)
        shutil.copy2(src_doc, os.path.join(case_dir, f"Case_{case_id}.docx"))
        shutil.copy2(src_txt, os.path.join(case_dir, "2-matrix.txt"))
        cases.append(case_dir)
    return cases